
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
import os
from PIL import ImageFile
from PIL import Image
import streamlit as st
ImageFile.LOAD_TRUNCATED_IMAGES = True
Image.MAX_IMAGE_PIXELS = None




def insert_images_to_word(image_paths, output_word_path):
    """
    将图像插入到 Word 文档中，每页一个图像。
    """
    doc = Document()
    #doc.add_heading("PDF to Word with Images", level=1)

    for i, image_path in enumerate(image_paths):
        #doc.add_paragraph(f"Page {i+1}:")
        doc.add_picture(image_path, width=Inches(6))  # 可调整宽度
        doc.add_paragraph("\n")  # 添加空行

    doc.save(output_word_path)
    st.write(f"Word document saved at: {output_word_path}")

def main():
    st.title('png转word工具')
    images_paths=[]#存图像路径
    uploaded_files = st.file_uploader("请选择文件:", accept_multiple_files=True, type=["PNG","JPG"])
    if st.button('图像传输完毕,点我开始生成word文件'):
        if uploaded_files:
            for uploaded_file in uploaded_files:
                images_paths.append(uploaded_file.name)
            #if 'text_input_active' not in st.session_state:
            st.session_state['text_input_active'] = True
            st.session_state['output_word_path'] = None
            st.session_state['image_paths']=images_paths
        else:
            st.write("尚未传输图像")

    if st.session_state.get('text_input_active', False):
        output_word_path = st.text_input(label='重命名你要生成word文档的名字', value='None')
        # 处理用户输入
        if output_word_path and output_word_path != 'None':
            st.session_state['output_word_path'] = output_word_path
            st.session_state['text_input_active'] = False  # 禁用输入框，表示流程完成
            # 显示用户的最终输入结果
    if st.session_state.get('output_word_path'):
        st.write(f"生成的文件名是: {st.session_state['output_word_path']}")
        st.write('您的输入是', output_word_path)
        output_word_path += ".docx"  # 输出 Word 文件路径
        st.write("Inserting images into Word document...")
        insert_images_to_word(st.session_state['image_paths'], output_word_path)
        del st.session_state['output_word_path']

    # output_word_path = "output.docx"  # 输出 Word 文件路径
    #
    # st.write("Inserting images into Word document...")
    # insert_images_to_word(images_paths, output_word_path)

main()

