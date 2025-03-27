

import os
import shutil
import tempfile
import zipfile
from pdf2image import convert_from_path
from PIL import Image
import streamlit as st
from docx import Document
from docx.shared import Inches


# 删除临时文件夹及其中内容的函数
def delete_temp_files(temp_dir):
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)


# 将 PDF 转换为图像并将图像插入到 Word 文档中
def pdf_to_word(pdf_path, temp_dir, word_filename):
    # 使用 pdf2image 将 PDF 转换为图像
    images = convert_from_path(pdf_path, dpi=300)  # 200 DPI 提供较好的图像质量

    # 创建 Word 文档
    doc = Document()

    # 设置每个页面的大小（例如 6x4 英寸）
    for i, image in enumerate(images):
        # 将图像保存到临时文件夹
        img_path = os.path.join(temp_dir, f"page_{i + 1}.png")
        image.save(img_path, "PNG")

        # 向 Word 文档添加一个新页面，插入图像
        doc.add_picture(img_path, width=Inches(6))  # 设置宽度为6英寸，自动按比例调整高度
        doc.add_paragraph("\n")  # 添加空行

    # 保存 Word 文档
    word_path = os.path.join(temp_dir, word_filename)
    doc.save(word_path)

    return word_path


# Streamlit 文件上传界面
def main():
    st.title("PDF 转 Word 工具")
    st.write("上传一个 PDF 文件，它将被转换成多个图像并存储在 Word 文档中。")

    # 上传 PDF 文件
    uploaded_file = st.file_uploader("选择一个 PDF 文件", type=["pdf"])

    # 输入自定义文件名前缀
    custom_filename = st.text_input("自定义输出 Word 文件名称（不包含扩展名）", "")

    # 只有当文件上传且转换未完成时才执行转换
    if uploaded_file and custom_filename:
        # 创建临时目录存放文件
        temp_dir = tempfile.mkdtemp()

        pdf_path = os.path.join(temp_dir, uploaded_file.name)

        # 将上传的 PDF 文件保存到本地临时目录
        with open(pdf_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # 使用 spinner 显示加载动画
        if st.button('开始转换'):
            with st.spinner('正在转换文件，请稍候...'):
                try:
                    # 将 PDF 转换为图像并插入到 Word 文件中
                    word_path = pdf_to_word(pdf_path, temp_dir, custom_filename + ".docx")
                    st.success("文件转换完成！")

                    # 提供下载按钮
                    with open(word_path, "rb") as word_file:
                        st.download_button("下载转换后的 Word 文件", word_file, file_name=custom_filename + ".docx")

                except Exception as e:
                    st.error(f"转换失败: {e}")

        # 清理临时文件夹和文件
        delete_temp_files(temp_dir)


main()
